import docx
from core.run_mapper import RunMapper

class DocxReader:
    def __init__(self, doc_path):
        self.doc_path = doc_path
        self.doc = docx.Document(doc_path)
        
    def extract_paragraphs(self):
        """
        Yields (paragraph_index, paragraph_obj, run_mapper)
        Includes paragraphs from body and tables.
        """
        index = 0
        
        # Extract from main document body
        for para in self.doc.paragraphs:
            if not para.text.strip():
                index += 1
                continue
            yield index, "body", para, RunMapper(para)
            index += 1
            
        # Extract from tables
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if not para.text.strip():
                            index += 1
                            continue
                        yield index, "table", para, RunMapper(para)
                        index += 1

    def clean_empty_paragraphs(self):
        """
        Removes consecutive empty paragraphs in the body to reduce multiple Enters to just one.
        Returns the number of removed paragraphs.
        """
        consecutive_empty = 0
        removed_count = 0
        
        # We only clean up the main body, tables usually have specific formatting
        for para in self.doc.paragraphs:
            if not para.text.strip():
                consecutive_empty += 1
                if consecutive_empty >= 2:
                    # Delete the paragraph element safely
                    p = para._element
                    if p.getparent() is not None:
                        p.getparent().remove(p)
                        para._p = para._element = None
                    removed_count += 1
            else:
                consecutive_empty = 0
                
        return removed_count

    def save(self, output_path):
        self.doc.save(output_path)
