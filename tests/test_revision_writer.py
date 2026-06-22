import pytest
from docx import Document
from core.revision_writer import RevisionWriter
from core.run_mapper import RunMapper
from docx.shared import RGBColor

def test_single_run_one_error():
    doc = Document()
    p = doc.add_paragraph()
    r1 = p.add_run("Tôi đang đi hoc.")
    
    mapper = RunMapper(p)
    writer = RevisionWriter(highlight_fallback=True)
    
    edits = [{
        'type': 'replace',
        'orig_start': 12,
        'orig_end': 15,
        'orig_text': 'hoc',
        'corr_text': 'học'
    }]
    
    results = writer.apply_edits(p, mapper, edits)
    assert len(results) == 1
    assert results[0]['applied'] is True
    assert results[0]['is_single_run'] is True
    
    # Check runs:
    # r1: "Tôi đang đi ", r2: "học", r3: "."
    assert len(p.runs) == 3
    assert p.runs[0].text == "Tôi đang đi "
    assert p.runs[1].text == "học"
    assert p.runs[1].font.color.rgb == RGBColor(255, 0, 0)
    assert p.runs[2].text == "."


def test_single_run_two_errors_reverse_order():
    doc = Document()
    p = doc.add_paragraph()
    r1 = p.add_run("Tôi đang di hoc nhe.")
    
    mapper = RunMapper(p)
    writer = RevisionWriter()
    
    edits = [
        {
            'type': 'replace',
            'orig_start': 9,
            'orig_end': 11,
            'orig_text': 'di',
            'corr_text': 'đi'
        },
        {
            'type': 'replace',
            'orig_start': 12,
            'orig_end': 15,
            'orig_text': 'hoc',
            'corr_text': 'học'
        }
    ]
    
    results = writer.apply_edits(p, mapper, edits)
    assert len(results) == 2
    for r in results:
        assert r['applied'] is True
        
    # Result should be: "Tôi đang " "đi" " " "học" " nhe."
    assert len(p.runs) == 5
    assert p.runs[0].text == "Tôi đang "
    assert p.runs[1].text == "đi"
    assert p.runs[2].text == " "
    assert p.runs[3].text == "học"
    assert p.runs[4].text == " nhe."


def test_multiple_runs_format_preservation():
    doc = Document()
    p = doc.add_paragraph()
    r1 = p.add_run("Hôm nay ")
    r2 = p.add_run("troi")
    r2.bold = True
    r3 = p.add_run(" đẹp.")
    
    mapper = RunMapper(p)
    writer = RevisionWriter()
    
    # "troi" is at index 8 to 12
    edits = [{
        'type': 'replace',
        'orig_start': 8,
        'orig_end': 12,
        'orig_text': 'troi',
        'corr_text': 'trời'
    }]
    
    results = writer.apply_edits(p, mapper, edits)
    assert results[0]['applied'] is True
    
    # Runs: "Hôm nay ", "trời" (bold), " đẹp."
    # Since 'troi' is exactly run 2, it splits into text_before="", corr="trời", text_after=""
    # Then text_before="" and text_after="" runs are dropped/empty, depending on implementation
    texts = [r.text for r in p.runs if r.text]
    assert "trời" in texts
    
    # Check formatting of the 'trời' run
    tr_run = [r for r in p.runs if r.text == "trời"][0]
    assert tr_run.bold is True
    assert tr_run.font.color.rgb == RGBColor(255, 0, 0)


def test_multi_run_edit_skip():
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("Thời ")
    r2 = p.add_run("tiêst")
    r2.bold = True
    p.add_run(" hôm nay.")
    
    mapper = RunMapper(p)
    writer = RevisionWriter()
    
    # original text: "Thời tiêst hôm nay."
    # Let's say model corrects "Thời tiêst" -> "Thời tiết"
    # orig_start: 0, orig_end: 10
    edits = [{
        'type': 'replace',
        'orig_start': 0,
        'orig_end': 10,
        'orig_text': 'Thời tiêst',
        'corr_text': 'Thời tiết'
    }]
    
    results = writer.apply_edits(p, mapper, edits)
    assert results[0]['applied'] is False
    assert results[0]['is_single_run'] is False
    assert results[0]['skip_reason'] == 'multi_run_edit_not_supported_yet'
    
    # Document shouldn't change
    assert p.runs[0].text == "Thời "
    assert p.runs[1].text == "tiêst"


def test_overlap_edit_skip():
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("Tôi đi hoc nhe.")
    
    mapper = RunMapper(p)
    writer = RevisionWriter()
    
    edits = [
        {
            'type': 'replace',
            'orig_start': 7,
            'orig_end': 10,
            'orig_text': 'hoc',
            'corr_text': 'học'
        },
        {
            'type': 'replace',
            'orig_start': 8,
            'orig_end': 10,
            'orig_text': 'oc',
            'corr_text': 'ọc'
        }
    ]
    
    results = writer.apply_edits(p, mapper, edits)
    # The first edit in the list is 7-10, the second is 8-10.
    # Reverse order sort: 8-10 first, then 7-10.
    # When 7-10 is processed, its end (10) > min_processed_start (8), so it overlaps!
    # Wait, reverse sort by orig_start: 
    # Edit 1: orig_start 8, orig_end 10. Processed. min_start = 8.
    # Edit 2: orig_start 7, orig_end 10. 10 > 8. Overlap!
    
    sorted_res = sorted(results, key=lambda x: x['orig_start'])
    
    assert sorted_res[1]['applied'] is True  # The 'oc' -> 'ọc' edit
    assert sorted_res[0]['applied'] is False # The 'hoc' -> 'học' edit
    assert sorted_res[0]['skip_reason'] == 'overlap'
